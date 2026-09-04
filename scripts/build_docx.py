"""
Renders a project markdown file into a Word document formatted to the
Section 3.1 standard of the BED 106 brief (Arial 12pt, 1-inch margins,
1.5 line spacing, numbered figure and table captions).

Handles the subset of markdown the report actually uses: ATX headings, pipe
tables, fenced code blocks, bullet and numbered lists, images, blockquotes,
horizontal rules, and inline **bold** / *italic* / `code`.

Run:  python3 scripts/build_docx.py            # builds every document
      python3 scripts/build_docx.py IN.md OUT.docx
"""

import hashlib
import io
import os
import re
import sys
import zipfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCUMENTS = [
    (os.path.join(ROOT, "reports", "Checkpoint_1_Report.md"),
     os.path.join(ROOT, "reports", "Checkpoint_1_Report.docx")),
    (os.path.join(ROOT, "docs", "checkpoint1_explained.md"),
     os.path.join(ROOT, "docs", "Checkpoint_1_Explained.docx")),
    (os.path.join(ROOT, "reports", "Checkpoint_2_Report.md"),
     os.path.join(ROOT, "reports", "Checkpoint_2_Report.docx")),
    (os.path.join(ROOT, "docs", "checkpoint2_explained.md"),
     os.path.join(ROOT, "docs", "Checkpoint_2_Explained.docx")),
]

INK = RGBColor(0x1F, 0x29, 0x33)
ACCENT = RGBColor(0x2F, 0x6F, 0x9F)
MUTED = RGBColor(0x7B, 0x89, 0x94)

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
IMAGE = re.compile(r"^!\[(.*?)\]\((.+?)\)$")


def add_runs(par, text):
    """Write text into a paragraph, honouring inline markdown emphasis."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = par.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = ACCENT
        elif piece.startswith("*") and piece.endswith("*"):
            par.add_run(piece[1:-1]).italic = True
        else:
            par.add_run(piece)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    for name, size, color in (("Heading 1", 16, ACCENT),
                              ("Heading 2", 14, INK),
                              ("Heading 3", 12, INK)):
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(14 if size > 12 else 10)
        st.paragraph_format.space_after = Pt(5)


def add_caption(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(10)
    par.paragraph_format.line_spacing = 1.0
    run = par.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def add_table(doc, rows, caption=None):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for cell, text in zip(table.rows[0].cells, header):
        cell.text = ""
        add_runs(cell.paragraphs[0], text)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row in body:
        cells = table.add_row().cells
        # Guard against a ragged row losing content off the end.
        for i, text in enumerate(row[:len(header)]):
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], text)
            para = cells[i].paragraphs[0]
            para.paragraph_format.line_spacing = 1.0
            for run in para.runs:
                run.font.size = Pt(9)
    if caption:
        add_caption(doc, caption)
    else:
        doc.add_paragraph()


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_code(doc, lines):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.25)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    run = par.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x3D, 0x47)


# A block ends at a blank line or at the start of any other markdown
# construct; everything before that is one soft-wrapped block.
BLOCK_START = re.compile(r"^(#|\||```|[-*]\s|\d{1,2}\.\s|>\s|!\[|---$)")


def gather(lines, i, first):
    """Join a block's soft-wrapped continuation lines into one string."""
    block = [first]
    while i < len(lines) and lines[i].strip() \
            and not BLOCK_START.match(lines[i].strip()):
        block.append(lines[i].strip())
        i += 1
    return " ".join(block), i


def convert(md, doc):
    lines = md.splitlines()
    i = 0
    counters = {"table": 0, "figure": 0}
    heading = ""
    pending_caption = None
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("Caption:"):
            pending_caption = stripped[len("Caption:"):].strip()
            i += 1
            continue

        if stripped == "---":
            doc.add_page_break()
            i += 1
            continue

        if stripped.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code(doc, block)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) \
                and set(lines[i + 1].strip()) <= set("|- :"):
            rows = [split_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            counters["table"] += 1
            label = pending_caption or heading
            pending_caption = None
            add_table(doc, rows, f"Table {counters['table']}. {label}".strip())
            continue

        m = IMAGE.match(stripped)
        if m:
            path = os.path.normpath(os.path.join(ROOT, "reports", m.group(2)))
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                counters["figure"] += 1
                label = pending_caption or m.group(1)
                pending_caption = None
                add_caption(doc, f"Figure {counters['figure']}. {label}".strip())
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading = stripped[level:].strip()
            doc.add_heading(heading, min(level, 3))
            i += 1
            continue

        m = re.match(r"^(?:[-*]\s+|\d{1,2}\.\s+)(.*)$", stripped)
        if m:
            style = "List Bullet" if stripped[0] in "-*" else "List Number"
            text, i = gather(lines, i + 1, m.group(1))
            add_runs(doc.add_paragraph(style=style), text)
            continue

        if stripped.startswith("> "):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.3)
            add_runs(par, stripped[2:])
            for run in par.runs:
                run.italic = True
            i += 1
            continue

        # Ordinary paragraph: join the soft-wrapped lines that follow it.
        text, i = gather(lines, i + 1, stripped)
        add_runs(doc.add_paragraph(), text)


def content_digest(data):
    """Hash a .docx by its internal parts, ignoring zip container framing.

    python-docx writes fresh timestamps and part ordering on every save, so
    two byte-different files can hold identical documents. Comparing the parts
    is what tells you whether anything actually changed.
    """
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        h = hashlib.sha256()
        for name in sorted(z.namelist()):
            h.update(name.encode())
            h.update(hashlib.sha256(z.read(name)).digest())
        return h.digest()


def build(src, out):
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(1.0)

    convert(open(src, encoding="utf-8").read(), doc)

    buf = io.BytesIO()
    doc.save(buf)
    new = buf.getvalue()

    # Leave the file alone when the document is unchanged, so rebuilding does
    # not show up as a modification in git.
    if os.path.exists(out):
        with open(out, "rb") as fh:
            if content_digest(fh.read()) == content_digest(new):
                print("unchanged", out)
                return

    with open(out, "wb") as fh:
        fh.write(new)
    print("wrote", out)


def main():
    if len(sys.argv) == 3:
        build(sys.argv[1], sys.argv[2])
        return
    for src, out in DOCUMENTS:
        build(src, out)


if __name__ == "__main__":
    main()
