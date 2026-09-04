"""
Builds the printable Individual Contribution Form (Form A).

Section 1.4 of the BED 106 brief requires a signed copy from every group
member with every checkpoint, so the document holds one page per role - four
ready-to-sign copies - followed by a reminder page of what each role actually
did in Checkpoints 1 and 2.

Formatted to the Section 3.1 standard: Arial 12pt, 1-inch margins.

Run:  python3 scripts/build_form_a.py
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "reports", "Form_A_Individual_Contribution.docx")

FONT = "Arial"
INK = RGBColor(0x1F, 0x29, 0x33)
ACCENT = RGBColor(0x2F, 0x6F, 0x9F)
MUTED = RGBColor(0x5A, 0x66, 0x72)

ROLES = ["Project Lead / Analyst", "Data Engineer",
         "Statistician / Modeler", "BI Developer / Visualizer"]

PROMPTS = {
    "Checkpoint 1 — Data Fundamentals & SQL Querying": [
        ("Project Lead / Analyst",
         "Business problem framing, the three key business questions, business "
         "interpretations of the query results, report assembly."),
        ("Data Engineer",
         "Dataset sourcing and documentation, data quality assessment, "
         "cleaning steps, ERD, schema creation, data load."),
        ("Statistician / Modeler",
         "Aggregate and trend queries, year-on-year and seasonality "
         "calculations, verification that reported figures match the data."),
        ("BI Developer / Visualizer",
         "Charts and figures, table formatting, report layout and "
         "presentation of results."),
    ],
    "Checkpoint 2 — Spreadsheet & Statistical Analysis": [
        ("Project Lead / Analyst",
         "Choice of variables for correlation and regression, written "
         "interpretations, report assembly, the Checkpoint 1 correction."),
        ("Data Engineer",
         "Export of the cleaned dataset into the workbook, named ranges, "
         "pivot cross-tabs and formula showcase, workbook self-checks."),
        ("Statistician / Modeler",
         "Descriptive statistics, frequency distribution, correlation "
         "coefficients, the regression and its significance test, the trend "
         "and seasonality analysis and its holdout check."),
        ("BI Developer / Visualizer",
         "Pivot charts, histogram, scatter plots with trendlines, forecast "
         "chart, workbook formatting and layout."),
    ],
}


def rule(paragraph):
    """Draw a writing line as a bottom border, so there is room to write."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "7B8794")
    borders.append(bottom)
    pPr.append(borders)


def field(doc, label, space_after=14):
    """A labelled write-in line."""
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(2)
    run = par.add_run(label)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    par.add_run("  ").font.size = Pt(11)
    rule(par)
    return par


def blank_line(doc, prefix=""):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(16)
    if prefix:
        run = par.add_run(prefix)
        run.font.name = FONT
        run.font.size = Pt(11)
        run.font.color.rgb = INK
    rule(par)
    return par


def heading(doc, text, size=15, color=ACCENT, after=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(after)
    run = par.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return par


def small(doc, text, italic=True, after=10):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(after)
    run = par.add_run(text)
    run.italic = italic
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    return par


def form_page(doc, role):
    heading(doc, "Individual Contribution Form", 16)
    heading(doc, "BED 106 — Business Analytics · Mini Capstone Project",
            11, INK, after=2)
    small(doc, "Talibon Polytechnic College · A.Y. 2026–2027, 1st Semester · "
               "Instructor: Jessie A. Melendres", after=16)

    field(doc, "Group Name / Number:")
    field(doc, "Checkpoint No.:")
    field(doc, "Member Name:")

    # The role is pre-filled so each printed page belongs to one member.
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(18)
    run = par.add_run("Role:  ")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    run = par.add_run(role)
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    rule(par)

    heading(doc, "Specific Contributions This Checkpoint", 12, INK, after=4)
    small(doc, "Write these in your own words. Be specific about what you "
               "personally did.", after=12)
    for n in (1, 2, 3):
        blank_line(doc, f"{n}.  ")
        blank_line(doc)

    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(12)
    par.paragraph_format.space_after = Pt(26)
    run = par.add_run("I certify that the contributions listed above are "
                      "accurate and reflect my actual participation.")
    run.font.name = FONT
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for i, (label, width) in enumerate((("Signature:", 3.9), ("Date:", 2.4))):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        par = cell.paragraphs[0]
        run = par.add_run(label)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(11)
        run.font.color.rgb = INK
        rule(par)


def prompts_page(doc):
    heading(doc, "What each role did", 15)
    small(doc, "Reminders only. Section 3.2 of the brief requires the work to "
               "be described in your own words, so use these to jog your "
               "memory rather than copying them.", after=14)

    for checkpoint, rows in PROMPTS.items():
        heading(doc, checkpoint, 12, INK, after=6)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, text in enumerate(("Role", "Work in this checkpoint")):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(text)
            run.bold = True
            run.font.name = FONT
            run.font.size = Pt(10)
        for role, work in rows:
            cells = table.add_row().cells
            for i, text in enumerate((role, work)):
                cells[i].text = ""
                run = cells[i].paragraphs[0].add_run(text)
                run.font.name = FONT
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True
        table.columns[0].width = Inches(1.9)
        table.columns[1].width = Inches(4.6)
        doc.add_paragraph().paragraph_format.space_after = Pt(10)


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(1.0)

    for i, role in enumerate(ROLES):
        if i:
            doc.add_page_break()
        form_page(doc, role)

    doc.add_page_break()
    prompts_page(doc)

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  {len(ROLES)} signable copies, one per role, plus a prompts page")


if __name__ == "__main__":
    main()
